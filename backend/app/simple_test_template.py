#!/usr/bin/env python3

from docx import Document
import os

def inspect_template():
    """Inspect the template file to see what's in it"""
    template_path = os.path.join(os.path.dirname(__file__), 'OfferLetterFinal.docx')
    
    if not os.path.exists(template_path):
        print(f"Template file not found: {template_path}")
        return
    
    try:
        doc = Document(template_path)
        print(f"Template loaded successfully. Number of paragraphs: {len(doc.paragraphs)}")
        
        # Check first few paragraphs
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f"Paragraph {i}: {para.text[:100]}...")
        
        # Check for any special characters or problematic content
        print("\nChecking for potential template syntax issues...")
        
    except Exception as e:
        print(f"Error loading template: {e}")

if __name__ == "__main__":
    inspect_template()
